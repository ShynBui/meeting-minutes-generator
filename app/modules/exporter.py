from inspect import signature

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.modules.schema import MeetingMinutes
from docx.shared import Pt, Cm
import os


def export_meeting_minutes_to_docx(meeting_minutes: MeetingMinutes, output_file: str) -> None:
    """
    Xuất meeting minutes thành file DOCX với định dạng chuẩn văn bản hành chính Việt Nam,
    mô phỏng layout như biểu mẫu (có phần tiêu ngữ, số hiệu, chữ ký, ...).

    Args:
        meeting_minutes (MeetingMinutes): Object chứa thông tin biên bản cuộc họp.
        output_file (str): Đường dẫn file DOCX đầu ra.
    """
    document = Document()

    # Thiết lập font mặc định (Times New Roman, size 13)
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # Đặt đơn vị đo lề (nếu cần)
    section = document.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # -------- PHẦN HEADER (TIÊU NGỮ, SỐ HIỆU) --------
    # Sử dụng bảng 2 cột để mô phỏng cột trái và cột phải
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    # Cột trái: Tên cơ quan, tổ chức
    cell_left_1 = table.cell(0, 0).paragraphs[0]
    cell_left_1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left_1 = cell_left_1.add_run("TÊN CƠ QUAN, TỔ CHỨC CHỦ QUẢN (1)\nTÊN CƠ QUAN, TỔ CHỨC (2)\n-------")
    run_left_1.font.name = 'Times New Roman'
    run_left_1.font.size = Pt(13)
    run_left_1.bold = True

    # Cột phải: Tiêu ngữ
    cell_right_1 = table.cell(0, 1).paragraphs[0]
    cell_right_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_right_1 = cell_right_1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n-------")
    run_right_1.font.name = 'Times New Roman'
    run_right_1.font.size = Pt(13)
    run_right_1.bold = True

    # Hàng 2: Số hiệu
    cell_left_2 = table.cell(1, 0).paragraphs[0]
    cell_left_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_left_2 = cell_left_2.add_run("")

    cell_right_2 = table.cell(1, 1).paragraphs[0]
    cell_right_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_right_2 = cell_right_2.add_run("Số: ....../BB-....(3)....")
    run_right_2.font.name = 'Times New Roman'
    run_right_2.font.size = Pt(13)

    # Xuống dòng sau bảng
    document.add_paragraph("")

    # -------- TIÊU ĐỀ CHÍNH --------
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_para.add_run("BIÊN BẢN CUỘC HỌP")
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(16)
    run_title.bold = True

    # -------- NỘI DUNG BIÊN BẢN --------
    document.add_paragraph("")

    # Thời gian bắt đầu (dùng gio_hop hoặc custom)
    # Ví dụ: "Thời gian bắt đầu: 10:00 - 11:30"
    if meeting_minutes.gio_hop:
        p_time = document.add_paragraph()
        p_time.add_run("Thời gian bắt đầu: ").bold = True
        p_time.add_run(meeting_minutes.gio_hop)

    # Địa điểm
    if meeting_minutes.dia_diem:
        p_place = document.add_paragraph()
        p_place.add_run("Địa điểm: ").bold = True
        p_place.add_run(meeting_minutes.dia_diem)

    # Thành phần
    # Gộp: Chủ trì, Thư ký, Thành viên tham dự
    document.add_paragraph("")

    document.add_paragraph("Thành phần cuộc họp: ", style="List Paragraph").runs[0].bold = True
    if meeting_minutes.chu_tri:
        document.add_paragraph(f"-Chủ trì: {meeting_minutes.chu_tri}", style="List Bullet")
    if meeting_minutes.nguoi_ghi_chep:
        document.add_paragraph(f"-Thư kí: {meeting_minutes.nguoi_ghi_chep}", style="List Bullet")
    if meeting_minutes.thanh_vien_tham_du:
        document.add_paragraph(f"- Các thành viên tham dự: " + ", ".join(meeting_minutes.thanh_vien_tham_du), style="List Bullet")

    # Mục tiêu cuộc họp
    if meeting_minutes.muc_tieu_cuoc_hop:
        doc_para = document.add_paragraph()
        doc_para.add_run("Nội dung, mục tiêu cuộc họp: ").bold = True
        doc_para.add_run(meeting_minutes.muc_tieu_cuoc_hop)

    document.add_paragraph("")

    # Chương trình nghị sự
    if meeting_minutes.chuong_trinh_nghi_su:
        doc_para = document.add_paragraph()
        doc_para.add_run("Chương trình nghị sự: ").bold = True
        for item in meeting_minutes.chuong_trinh_nghi_su:
            document.add_paragraph(item, 'List Bullet')
        document.add_paragraph("")


    # Nội dung thảo luận
    if meeting_minutes.noi_dung_thao_luan:
        doc_para = document.add_paragraph()
        doc_para.add_run("Nội dung thảo luận: ").bold = True
        for topic, points in meeting_minutes.noi_dung_thao_luan.items():
            sub_heading = document.add_paragraph(topic)
            sub_heading.style = 'List Bullet'
            for point in points:
                document.add_paragraph(point, 'List Bullet 2')

    document.add_paragraph("")

    # Các quyết định
    if meeting_minutes.cac_quyet_dinh:
        doc_para = document.add_paragraph()
        doc_para.add_run("Các quyết định:").bold = True
        for decision in meeting_minutes.cac_quyet_dinh:
            document.add_paragraph(decision, style ='List Bullet')

    document.add_paragraph("")


    # Kết luận
    if meeting_minutes.ket_luan:
        doc_para = document.add_paragraph()
        doc_para.add_run("Kết luận: ").bold = True
        doc_para.add_run(meeting_minutes.ket_luan)

    document.add_paragraph("")

    # Tài liệu đính kèm
    if meeting_minutes.tai_lieu_dinh_kem:
        doc_para = document.add_paragraph()
        doc_para.add_run("Tài liệu đính kèm: ").bold = True
        for item in meeting_minutes.tai_lieu_dinh_kem:
            document.add_paragraph(item, 'List Bullet')
    else:
        document.add_paragraph("Tài liệu đính kèm: Không có")

    # Ghi chú

    document.add_paragraph("")
    if meeting_minutes.ghi_chu:
        doc_para = document.add_paragraph()
        doc_para.add_run("Ghi chú: ").bold = True
        doc_para.add_run(meeting_minutes.ghi_chu)
    else:
        document.add_paragraph("Ghi chú: Không có")

    document.add_paragraph("")


    # -------- PHẦN CHỮ KÝ --------
    # Tạo bảng 2 cột cho chữ ký (Chủ trì - Thư ký)
    signature_table = document.add_table(rows=1, cols=2)
    signature_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature_table.columns[0].width = Cm(8)
    signature_table.columns[1].width = Cm(8)

    cell_left = signature_table.cell(0, 0).paragraphs[0]
    cell_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_left = cell_left.add_run("CHỦ TRÌ\n\n\n\n\n(Ký, ghi rõ họ tên)")
    run_left.font.name = "Times New Roman"
    run_left.font.size = Pt(13)

    cell_right = signature_table.cell(0, 1).paragraphs[0]
    cell_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_right = cell_right.add_run("THƯ KÝ\n\n\n\n\n(Ký, ghi rõ họ tên)")
    run_right.font.name = "Times New Roman"
    run_right.font.size = Pt(13)

    # Lưu file DOCX
    document.save(output_file)


# --- Phần test chạy độc lập ---
if __name__ == "__main__":
    # Giả sử bạn có một đối tượng MeetingMinutes mẫu
    sample_minutes = MeetingMinutes(
        ngay_hop="30/03/2025",
        gio_hop="10:00 - 11:30",
        dia_diem="Phòng họp A",
        chu_tri="Nguyễn Văn A",
        nguoi_ghi_chep="Trần Thị B",
        thanh_vien_tham_du=["Nguyễn Văn A", "Trần Thị B", "Lê Văn C", "Phạm Thị D"],
        muc_tieu_cuoc_hop="Đánh giá tiến độ và lên kế hoạch giai đoạn tiếp theo",
        chuong_trinh_nghi_su=["Báo cáo tiến độ", "Phân tích khó khăn", "Giải pháp và phân công nhiệm vụ"],
        noi_dung_thao_luan={
            "Báo cáo tiến độ": ["Đã hoàn thành 70% công việc"],
            "Khó khăn": ["Thiếu nhân sự cho giai đoạn thử nghiệm"],
            "Giải pháp": ["Đề xuất tăng thêm 2 tester", "Lê Văn C phụ trách", "Hạn chót: 15/04/2025"]
        },
        cac_quyet_dinh=["Phê duyệt bổ sung nhân sự", "Tái phân công nhiệm vụ"],
        ket_luan="Thống nhất hạn hoàn thành và họp tiếp theo vào 20/04/2025",
        tai_lieu_dinh_kem=None,
        ghi_chu="Cần chú ý việc phân bổ nhiệm vụ và theo dõi tiến độ một cách chặt chẽ."
    )

    print("----- Meeting Minutes gốc -----")
    print(sample_minutes.model_dump())

    # Xuất file DOCX từ meeting minutes đã được chỉnh sửa
    output_docx = "meeting_minutes_refined.docx"
    export_meeting_minutes_to_docx(sample_minutes, output_docx)
    print(f"Meeting minutes đã được xuất ra file: {output_docx}")